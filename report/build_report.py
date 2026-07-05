"""
build_report.py
---------------
Erzeugt den schriftlichen Projektbericht als Word-Datei (.docx, ca. 6-8
Seiten) aus den Ergebnissen der Analyse-Pipeline. Liest outputs/metrics.json
und bindet die Abbildungen aus outputs/figures/ ein.

Voraussetzung: src/main.py wurde vorher ausgefuehrt, sodass metrics.json und
die Abbildungen existieren.

Aufruf:
    python report/build_report.py
"""

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parent.parent
FIG = PROJECT_ROOT / "outputs" / "figures"
METRICS_PATH = PROJECT_ROOT / "outputs" / "metrics.json"
OUTPUT_DOCX = PROJECT_ROOT / "report" / "Projektbericht.docx"

NAVY = RGBColor(0x1F, 0x3B, 0x5C)
GREY = RGBColor(0x5A, 0x5A, 0x5A)


# --------------------------------------------------------------------------
# Hilfsfunktionen
# --------------------------------------------------------------------------
def fmt(x):
    return f"{x:.4f}".replace(".", ",")


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def add_figure(doc, filename, caption, width_inches=5.8):
    path = FIG / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    for run in cap.runs:
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = GREY


def add_toc(doc):
    """Fuegt ein automatisches Inhaltsverzeichnis-Feld ein (aktualisiert
    sich beim Oeffnen in Word/LibreOffice ueber F9)."""
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-2" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "Rechtsklick hier → Feld aktualisieren, um das Inhaltsverzeichnis zu erzeugen."
    fldChar2.append(t)
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_kpi_table(doc, header, rows):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, text in enumerate(header):
        set_cell_shading(hdr[i], "1F3B5C")
        para = hdr[i].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for r, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            if r % 2 == 1:
                set_cell_shading(cells[i], "F0F3F7")
            para = cells[i].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(str(val))
            run.font.size = Pt(10)
            if i == 0:
                run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def style_document(doc):
    """Setzt Standard-Schrift und Ueberschriften-Farben."""
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for level, size in [("Heading 1", 16), ("Heading 2", 13)]:
        st = doc.styles[level]
        st.font.color.rgb = NAVY
        st.font.size = Pt(size)
        st.font.name = "Calibri"


# --------------------------------------------------------------------------
# Bericht
# --------------------------------------------------------------------------
def build():
    metrics = json.loads(METRICS_PATH.read_text())
    ds = metrics["dataset"]
    split = metrics["split"]
    tm = metrics["test_metrics"]
    cv = metrics["cross_validation_f1"]
    lr = tm["Logistic Regression (Baseline)"]
    rf = tm["Random Forest (getunt)"]
    lr_cv = cv["Logistic Regression (Baseline)"]
    rf_cv = cv["Random Forest (getunt)"]
    best = metrics["best_model"]
    rf_params = metrics["best_params_random_forest"]
    top_corr = metrics["top_target_correlations"]
    top_imp = metrics["top_feature_importances"]

    doc = Document()
    style_document(doc)

    # ---------------- Titelseite ----------------
    for _ in range(3):
        doc.add_paragraph()
    sub = doc.add_paragraph("Projektarbeit – Maschinelles Lernen sowie "
                            "Datenstatistik und -analyse")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(12)
    sub.runs[0].font.color.rgb = GREY
    sub.runs[0].bold = True

    doc.add_paragraph()
    title = doc.add_paragraph("Klassifikation von Brusttumoren")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(30)
    title.runs[0].bold = True
    title.runs[0].font.color.rgb = NAVY

    sub2 = doc.add_paragraph("Analyse des Wisconsin Diagnostic Breast Cancer "
                             "(WDBC) Datensatzes")
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2.runs[0].font.size = Pt(16)
    sub2.runs[0].bold = True
    sub2.runs[0].font.color.rgb = NAVY

    for _ in range(2):
        doc.add_paragraph()
    lead = doc.add_paragraph(
        "Datengestützte Diagnose gutartiger und bösartiger Tumoren mit Methoden "
        "der deskriptiven Statistik und des überwachten maschinellen Lernens")
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead.runs[0].font.size = Pt(12)

    for _ in range(6):
        doc.add_paragraph()
    meta = doc.add_paragraph(
        f"Datensatz: {ds['n_samples']} Fälle · {ds['n_features']} Merkmale · "
        f"Zielklassen: gutartig / bösartig\n"
        f"Modelle: Logistic Regression (Baseline) und Random Forest (Hauptmodell)\n"
        f"Bewertungsmetrik im Fokus: F1-Score")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in meta.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = GREY

    doc.add_page_break()

    # ---------------- Inhaltsverzeichnis ----------------
    h = doc.add_paragraph("Inhaltsverzeichnis")
    h.runs[0].bold = True
    h.runs[0].font.size = Pt(14)
    h.runs[0].font.color.rgb = NAVY
    add_toc(doc)
    doc.add_page_break()

    # ---------------- 1 Einleitung ----------------
    doc.add_heading("1  Einleitung und Zielsetzung", level=1)
    add_body(doc,
        "Brustkrebs ist eine der häufigsten Krebserkrankungen weltweit. Eine "
        "frühzeitige und zuverlässige Unterscheidung zwischen gutartigen "
        "(benignen) und bösartigen (malignen) Tumoren ist entscheidend für die "
        "Prognose der Patientinnen und Patienten. In der klinischen Praxis wird "
        "dazu häufig eine Feinnadelaspiration (FNA) durchgeführt: Aus einem "
        "Tumor werden Zellen entnommen, digital fotografiert und die Zellkerne "
        "vermessen. Aus diesen Vermessungen lassen sich numerische Merkmale "
        "ableiten, die einen Rückschluss auf die Bösartigkeit erlauben.")
    add_body(doc,
        "Ziel dieses Projekts ist es, den Wisconsin Diagnostic Breast Cancer "
        "(WDBC) Datensatz mit Methoden der Datenstatistik und des maschinellen "
        "Lernens zu analysieren und ein Modell zu entwickeln, das die Diagnose "
        "möglichst zuverlässig aus den gemessenen Zellkern-Merkmalen vorhersagt. "
        "Das Projekt verbindet damit beide Fächer: Der statistische Teil "
        "beschreibt und visualisiert die Daten, der Machine-Learning-Teil "
        "trainiert und bewertet ein Vorhersagemodell.")
    doc.add_heading("1.1  Leitfragen", level=2)
    add_bullet(doc, "Wie sind die Daten verteilt, und welche Merkmale "
                    "unterscheiden gutartige von bösartigen Tumoren am stärksten?")
    add_bullet(doc, "Wie stark sind die Merkmale untereinander korreliert, und "
                    "was bedeutet das für die Modellierung?")
    add_bullet(doc, "Wie gut lässt sich die Diagnose vorhersagen, gemessen "
                    "insbesondere am F1-Score?")
    add_bullet(doc, "Welche Fehler macht das Modell, und wie sind diese im "
                    "medizinischen Kontext zu bewerten?")

    # ---------------- 2 Datensatz ----------------
    doc.add_heading("2  Der Datensatz", level=1)
    add_body(doc,
        f"Der WDBC-Datensatz wurde an der University of Wisconsin erhoben "
        f"(Wolberg, Street und Mangasarian, 1995) und ist ein etablierter "
        f"Benchmark-Datensatz. Er umfasst {ds['n_samples']} Fälle mit je "
        f"{ds['n_features']} numerischen Merkmalen sowie der Diagnose als "
        f"Zielvariable. Es gibt keine fehlenden Werte. Von {ds['n_benign']} "
        f"Fällen ist der Tumor gutartig (B), von {ds['n_malignant']} bösartig "
        f"(M) – die Klassen sind also leicht unausgewogen (rund "
        f"{100*ds['n_benign']/ds['n_samples']:.0f} % zu "
        f"{100*ds['n_malignant']/ds['n_samples']:.0f} %).")
    doc.add_heading("2.1  Merkmalsaufbau", level=2)
    add_body(doc,
        "Für jeden Zellkern werden zehn Basismerkmale berechnet: radius, "
        "texture, perimeter, area, smoothness, compactness, concavity, concave "
        "points, symmetry und fractal dimension. Von jedem dieser Merkmale "
        "liegen drei Ausprägungen vor:")
    add_bullet(doc, "mean – der Mittelwert über alle Zellkerne eines Bildes,")
    add_bullet(doc, "se – der Standardfehler (Streuung der Messung),")
    add_bullet(doc, "worst – der Mittelwert der drei größten Werte.")
    add_body(doc,
        "So ergeben sich 10 × 3 = 30 Merkmale. Die ID-Spalte wird verworfen, da "
        "sie ein reiner Identifikator ohne Vorhersageinformation ist. Die "
        "Zielvariable wird binär kodiert, wobei bösartig (M) als positive Klasse "
        "(1) definiert wird – das ist der medizinisch kritische Fall, den das "
        "Modell erkennen soll.")
    add_figure(doc, "01_class_distribution.png",
        "Abbildung 1: Klassenverteilung. Die gutartigen Fälle sind in der "
        "Überzahl, weshalb die reine Trefferquote (Accuracy) als alleinige "
        "Metrik irreführend sein kann.", width_inches=4.3)

    # ---------------- 3 Methodik ----------------
    doc.add_heading("3  Methodik", level=1)
    doc.add_heading("3.1  Statistische Analyse", level=2)
    add_body(doc,
        "Zunächst wird eine explorative Datenanalyse (EDA) durchgeführt. Dabei "
        "werden deskriptive Kennzahlen (Mittelwert, Median, Standardabweichung, "
        "Minimum, Maximum) berechnet und die Verteilungen der Merkmale getrennt "
        "nach Diagnose visualisiert. Über die Pearson-Korrelation werden sowohl "
        "die Zusammenhänge der Merkmale untereinander als auch ihr Zusammenhang "
        "mit der Zielvariable quantifiziert.")
    doc.add_heading("3.2  Maschinelles Lernen", level=2)
    add_body(doc,
        f"Der Datensatz wird stratifiziert im Verhältnis 80/20 in Trainings- "
        f"und Testdaten aufgeteilt ({split['n_train']} zu {split['n_test']} "
        f"Fälle). Stratifiziert bedeutet, dass das Klassenverhältnis in beiden "
        f"Teilmengen erhalten bleibt. Alle Modelle sind in einer sklearn-"
        f"Pipeline mit vorgeschalteter Standardisierung (StandardScaler) "
        f"gekapselt. Dadurch wird der Skalierer ausschließlich auf den "
        f"Trainingsdaten angepasst; ein Informationsabfluss aus den Testdaten "
        f"(Data Leakage) wird vermieden.")
    add_body(doc,
        "Als Baseline dient eine Logistische Regression. Das Hauptmodell ist ein "
        "Random Forest – ein Ensemble aus vielen Entscheidungsbäumen, das robust "
        "gegenüber Ausreißern ist und interpretierbare Merkmalswichtigkeiten "
        "liefert. Die Hyperparameter des Random Forest werden per Gittersuche "
        "(GridSearchCV) mit 5-facher stratifizierter Kreuzvalidierung optimiert, "
        "wobei direkt auf den F1-Score optimiert wird.")
    doc.add_heading("3.3  Bewertungsmetriken und der F1-Score", level=2)
    add_body(doc,
        "Da die Klassen unausgewogen sind, ist die Accuracy allein nicht "
        "aussagekräftig. Im Mittelpunkt steht daher der F1-Score. Er ist das "
        "harmonische Mittel aus Precision und Recall:")
    formula = doc.add_paragraph("F1 = 2 · (Precision · Recall) / (Precision + Recall)")
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula.runs[0].italic = True
    add_body(doc,
        "Precision (Genauigkeit) beantwortet: Wie viele der als bösartig "
        "vorhergesagten Fälle sind tatsächlich bösartig? Recall (Sensitivität) "
        "beantwortet: Wie viele der tatsächlich bösartigen Fälle wurden erkannt? "
        "Der F1-Score ist genau dann hoch, wenn beide Werte hoch sind – er "
        "bestraft ein Modell, das eine der beiden Größen vernachlässigt. "
        "Ergänzend werden Accuracy, die Konfusionsmatrix und die ROC-AUC "
        "(Fläche unter der ROC-Kurve) betrachtet.")

    # ---------------- 4 EDA ----------------
    doc.add_heading("4  Ergebnisse der explorativen Datenanalyse", level=1)
    add_body(doc,
        "Die Verteilungen der Merkmale unterscheiden sich deutlich zwischen den "
        "beiden Diagnosen. Bösartige Tumoren weisen im Schnitt größere Radien, "
        "Umfänge, Flächen und ausgeprägtere Konkavitäten auf. Abbildung 2 zeigt "
        "dies exemplarisch: Die roten Verteilungen (bösartig) sind gegenüber den "
        "blauen (gutartig) nach rechts verschoben. Gleichzeitig überlappen sich "
        "die Verteilungen, sodass kein einzelnes Merkmal die Klassen perfekt "
        "trennt – erst die Kombination mehrerer Merkmale ermöglicht eine "
        "zuverlässige Vorhersage.")
    add_figure(doc, "02_feature_distributions.png",
        "Abbildung 2: Dichteverteilungen (KDE) ausgewählter Merkmale, getrennt "
        "nach Diagnose. Bösartige Fälle sind systematisch zu größeren Werten "
        "verschoben.", width_inches=6.2)
    add_figure(doc, "03_boxplots.png",
        "Abbildung 3: Boxplots verdeutlichen die Verschiebung der Mediane und "
        "die größere Streuung bei bösartigen Tumoren.", width_inches=6.2)

    doc.add_heading("4.1  Korrelation mit der Diagnose", level=2)
    first_feat = list(top_corr.items())[0]
    add_body(doc,
        f"Die stärksten Zusammenhänge mit der Diagnose zeigen die worst- und "
        f"mean-Varianten von concave points, perimeter, radius und area. Das "
        f"Merkmal mit der höchsten Korrelation ist „{first_feat[0]}“ "
        f"(r = {first_feat[1]:.2f}). Alle Top-Merkmale sind positiv korreliert: "
        f"Größere Werte gehen mit einer höheren Wahrscheinlichkeit für "
        f"Bösartigkeit einher.")
    add_figure(doc, "05_target_correlations.png",
        "Abbildung 4: Merkmale mit dem stärksten Zusammenhang zur Diagnose "
        "(Pearson-Korrelation).", width_inches=5.0)

    doc.add_heading("4.2  Korrelation der Merkmale untereinander", level=2)
    add_body(doc,
        "Viele Merkmale sind untereinander stark korreliert – etwa radius, "
        "perimeter und area, die geometrisch direkt zusammenhängen. Solche "
        "Multikollinearität kann lineare Modelle destabilisieren, stört einen "
        "Random Forest jedoch kaum. Sie bedeutet auch, dass die effektive "
        "Informationsmenge geringer ist, als die 30 Merkmale vermuten lassen.")
    add_figure(doc, "04_correlation_heatmap.png",
        "Abbildung 5: Korrelationsmatrix aller 30 Merkmale. Die hellroten Blöcke "
        "zeigen Gruppen stark korrelierter Merkmale.", width_inches=6.0)

    doc.add_heading("4.3  Trennbarkeit der Klassen (PCA)", level=2)
    add_body(doc,
        "Eine Hauptkomponentenanalyse (PCA) projiziert die 30 Merkmale auf zwei "
        "Dimensionen. Bereits in dieser stark reduzierten Darstellung sind die "
        "beiden Klassen weitgehend getrennt (Abbildung 6). Das bestätigt die in "
        "der Literatur beschriebene nahezu lineare Trennbarkeit des Datensatzes "
        "und erklärt, warum schon einfache Modelle sehr gute Ergebnisse "
        "erzielen.")
    add_figure(doc, "06_pca_scatter.png",
        "Abbildung 6: PCA-Projektion. Die beiden Klassen sind bereits in zwei "
        "Dimensionen weitgehend trennbar.", width_inches=4.8)

    # ---------------- 5 Modellierung ----------------
    doc.add_heading("5  Modellierung und Ergebnisse", level=1)
    doc.add_heading("5.1  Kreuzvalidierung auf den Trainingsdaten", level=2)
    add_body(doc,
        f"Auf den Trainingsdaten wird jedes Modell per 5-facher Kreuzvalidierung "
        f"bewertet. Die Logistische Regression erreicht einen mittleren F1-Score "
        f"von {fmt(lr_cv['mean'])} (± {fmt(lr_cv['std'])}), der getunte Random "
        f"Forest {fmt(rf_cv['mean'])} (± {fmt(rf_cv['std'])}). Beide Modelle "
        f"sind also bereits vor dem Testset sehr treffsicher und stabil. Die als "
        f"optimal ermittelten Hyperparameter des Random Forest sind: "
        f"n_estimators = {rf_params['n_estimators']}, "
        f"max_depth = {rf_params['max_depth']}, "
        f"min_samples_leaf = {rf_params['min_samples_leaf']}, "
        f"max_features = „{rf_params['max_features']}“.")
    add_figure(doc, "10_cv_comparison.png",
        "Abbildung 7: Mittlerer F1-Score der Kreuzvalidierung mit "
        "Standardabweichung als Fehlerbalken.", width_inches=4.8)

    doc.add_heading("5.2  Bewertung auf dem Testset", level=2)
    add_body(doc,
        f"Die entscheidende Bewertung erfolgt auf dem zuvor unangetasteten "
        f"Testset ({split['n_test']} Fälle). Die folgende Tabelle fasst die "
        f"zentralen Kennzahlen beider Modelle zusammen.")
    add_kpi_table(doc,
        ["Metrik", "Log. Regression", "Random Forest"],
        [
            ["Accuracy", fmt(lr["accuracy"]), fmt(rf["accuracy"])],
            ["Precision", fmt(lr["precision"]), fmt(rf["precision"])],
            ["Recall (Sensitivität)", fmt(lr["recall"]), fmt(rf["recall"])],
            ["F1-Score", fmt(lr["f1"]), fmt(rf["f1"])],
            ["ROC-AUC", fmt(lr["roc_auc"]), fmt(rf["roc_auc"])],
        ])
    add_body(doc,
        f"Beide Modelle liefern sehr gute Ergebnisse mit F1-Scores um 0,94. "
        f"Bemerkenswert ist, dass die einfache Logistische Regression den Random "
        f"Forest beim F1-Score in diesem Testsplit knapp übertrifft "
        f"({fmt(lr['f1'])} gegenüber {fmt(rf['f1'])}). Das ist kein Zufall, "
        f"sondern eine direkte Folge der nahezu linearen Trennbarkeit des "
        f"Datensatzes: Ein lineares Modell passt hier hervorragend, während die "
        f"zusätzliche Flexibilität des Random Forest keinen Vorteil bringt. Der "
        f"Random Forest erzielt dafür eine perfekte Precision von "
        f"{fmt(rf['precision'])} – jede seiner Bösartig-Vorhersagen ist korrekt, "
        f"er übersieht aber mehr tatsächlich bösartige Fälle.")

    doc.add_heading("5.3  Konfusionsmatrix und Fehleranalyse", level=2)
    lr_cm = lr["confusion_matrix"]
    rf_cm = rf["confusion_matrix"]
    add_body(doc,
        f"Die Konfusionsmatrix des besten Modells ({best}) zeigt die Verteilung "
        f"der Fehler im Detail. Von {split['n_test']} Testfällen werden nur "
        f"wenige falsch klassifiziert. Medizinisch besonders relevant sind "
        f"falsch-negative Fälle: ein bösartiger Tumor, der fälschlich als "
        f"gutartig eingestuft wird. Die Logistische Regression produziert "
        f"{lr_cm[1][0]} solcher falsch-negativen und {lr_cm[0][1]} "
        f"falsch-positive Fälle. Der Random Forest hat mit {rf_cm[1][0]} "
        f"falsch-negativen Fällen hier mehr Übersehungen, dafür keinen einzigen "
        f"falsch-positiven.")
    add_figure(doc, "07_confusion_matrix.png",
        "Abbildung 8: Konfusionsmatrix des besten Modells. Die Nebendiagonale "
        "enthält die Fehlklassifikationen.", width_inches=4.3)
    add_figure(doc, "08_roc_curves.png",
        "Abbildung 9: ROC-Kurven beider Modelle. Beide liegen nahe der linken "
        "oberen Ecke (AUC > 0,99), also nahezu perfekt.", width_inches=4.7)

    doc.add_heading("5.4  Wichtigste Merkmale", level=2)
    imp_names = ", ".join(d["feature"] for d in top_imp[:5])
    add_body(doc,
        f"Der Random Forest liefert eine Rangfolge der Merkmalswichtigkeiten. "
        f"Die fünf einflussreichsten Merkmale sind: {imp_names}. Diese decken "
        f"sich weitgehend mit den Merkmalen, die schon in der EDA die stärkste "
        f"Korrelation zur Diagnose zeigten – Größen- und Formmerkmale der "
        f"Zellkerne (perimeter, area, radius, concave points) sind für die "
        f"Diagnose am aussagekräftigsten. Das Modell bestätigt damit die "
        f"statistische Analyse und ist zugleich medizinisch plausibel.")
    add_figure(doc, "09_feature_importance.png",
        "Abbildung 10: Die 15 wichtigsten Merkmale nach Gini-Importance des "
        "Random Forest.", width_inches=5.6)

    # ---------------- 6 Diskussion ----------------
    doc.add_heading("6  Diskussion", level=1)
    add_body(doc,
        "Die Ergebnisse zeigen, dass sich die Diagnose aus den Zellkern-"
        "Merkmalen sehr zuverlässig vorhersagen lässt. Mehrere Punkte verdienen "
        "eine kritische Einordnung.")
    doc.add_heading("6.1  Einfaches Modell schlägt komplexes Modell", level=2)
    add_body(doc,
        "Dass die Logistische Regression den Random Forest knapp übertrifft, ist "
        "ein lehrreiches Ergebnis. Es widerlegt die verbreitete Annahme, "
        "komplexere Modelle seien grundsätzlich besser. Bei nahezu linear "
        "trennbaren Daten mit begrenztem Stichprobenumfang ist ein einfaches, "
        "gut regularisiertes Modell oft die bessere Wahl – es ist "
        "interpretierbarer, schneller und weniger anfällig für Overfitting.")
    doc.add_heading("6.2  Precision-Recall-Abwägung im medizinischen Kontext", level=2)
    add_body(doc,
        "Die Wahl des Modells hängt vom Anwendungsziel ab. Im Screening ist ein "
        "hoher Recall entscheidend, weil ein übersehener bösartiger Tumor "
        "(falsch-negativ) schwerwiegende Folgen hat. Die Logistische Regression "
        "mit höherem Recall wäre hier vorzuziehen. Der Random Forest mit "
        "perfekter Precision wäre dagegen dort sinnvoll, wo jede Positiv-Meldung "
        "teure Folgeuntersuchungen auslöst. In der Praxis würde man die "
        "Entscheidungsschwelle bewusst so einstellen, dass falsch-negative Fälle "
        "minimiert werden.")
    doc.add_heading("6.3  Grenzen der Aussagekraft", level=2)
    add_bullet(doc, "Der Datensatz ist mit 569 Fällen vergleichsweise klein; "
                    "Ergebnisse schwanken je nach Zufallssplit.")
    add_bullet(doc, "Die Daten stammen aus einer einzigen Klinik der 1990er "
                    "Jahre und sind nicht ohne Weiteres auf andere Populationen "
                    "oder Geräte übertragbar.")
    add_bullet(doc, "Ein reales Diagnosesystem müsste klinisch validiert und auf "
                    "unabhängigen, aktuellen Daten geprüft werden; es ersetzt "
                    "keine ärztliche Beurteilung.")

    # ---------------- 7 Fazit ----------------
    doc.add_heading("7  Fazit und Ausblick", level=1)
    add_body(doc,
        f"In diesem Projekt wurde der WDBC-Datensatz vollständig analysiert – von "
        f"der deskriptiven Statistik über die Visualisierung bis zur überwachten "
        f"Klassifikation. Die explorative Analyse zeigte klare, medizinisch "
        f"plausible Unterschiede zwischen gutartigen und bösartigen Tumoren, "
        f"insbesondere bei Größen- und Formmerkmalen der Zellkerne. Beide "
        f"trainierten Modelle erreichten auf dem Testset einen F1-Score um 0,94 "
        f"und eine ROC-AUC von über 0,99. Das beste Modell war die Logistische "
        f"Regression mit einem F1-Score von {fmt(lr['f1'])}.")
    add_body(doc,
        "Für die Zukunft bieten sich mehrere Erweiterungen an: der Vergleich "
        "weiterer Modelle (z. B. Support Vector Machines oder Gradient "
        "Boosting), eine gezielte Merkmalsauswahl zur Reduktion der "
        "Multikollinearität, die Optimierung der Entscheidungsschwelle auf "
        "maximalen Recall sowie eine Kalibrierung der vorhergesagten "
        "Wahrscheinlichkeiten. Insgesamt zeigt das Projekt, wie statistische "
        "Analyse und maschinelles Lernen zusammenwirken, um aus Messdaten eine "
        "belastbare, nachvollziehbare Entscheidungsunterstützung zu gewinnen.")

    # ---------------- 8 Literatur ----------------
    doc.add_heading("8  Literatur und Werkzeuge", level=1)
    add_bullet(doc, "Wolberg, W. H., Street, W. N., Mangasarian, O. L. (1995): "
                    "Wisconsin Diagnostic Breast Cancer (WDBC). UCI Machine "
                    "Learning Repository.")
    add_bullet(doc, "Street, W. N., Wolberg, W. H., Mangasarian, O. L. (1993): "
                    "Nuclear feature extraction for breast tumor diagnosis. "
                    "IS&T/SPIE Symposium on Electronic Imaging.")
    add_bullet(doc, "Pedregosa et al. (2011): Scikit-learn: Machine Learning in "
                    "Python. Journal of Machine Learning Research 12.")
    add_bullet(doc, "Verwendete Werkzeuge: Python 3.11, pandas, NumPy, "
                    "scikit-learn, matplotlib, seaborn.")

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_DOCX))
    print(f"Bericht gespeichert: {OUTPUT_DOCX.relative_to(PROJECT_ROOT)}")


if __name__ == "__main__":
    build()
